"""
One-time helper to create beginner-friendly sample PDF and DOCX documents.

Run:
    python scripts/create_sample_documents.py
"""

from pathlib import Path

from docx import Document as DocxDocument
from fpdf import FPDF


OUTPUT_DIR = Path("data/documents")


class SimplePDF(FPDF):
    """Tiny helper so multi-page sample PDFs stay readable."""

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", size=9)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def add_heading(pdf: FPDF, text: str, size: int = 14) -> None:
    """Write a bold section heading."""
    pdf.set_font("Helvetica", "B", size)
    pdf.cell(0, 8, text)
    pdf.ln(10)


def add_wrapped_text(pdf: FPDF, text: str, size: int = 12) -> None:
    """Write a paragraph with basic wrapping."""
    pdf.set_font("Helvetica", size=size)
    pdf.multi_cell(0, 6, text)
    pdf.ln(2)


def create_company_handbook_pdf(path: Path) -> None:
    """Create a sample employee handbook PDF."""
    pdf = SimplePDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.add_page()
    add_heading(pdf, "Aspire Demo Company Handbook", size=18)
    add_wrapped_text(
        pdf,
        "Welcome to Aspire Demo Company. This handbook explains working hours, "
        "leave policy, remote work, and employee benefits for all full-time staff.",
    )

    add_heading(pdf, "1. Working Hours")
    add_wrapped_text(
        pdf,
        "Standard working hours are Monday to Friday, 9:00 AM to 6:00 PM, "
        "including a one-hour lunch break. Core collaboration hours are "
        "10:00 AM to 4:00 PM. Employees may request flexible start times "
        "between 8:00 AM and 10:00 AM with manager approval.",
    )

    add_heading(pdf, "2. Leave Policy")
    add_wrapped_text(
        pdf,
        "Full-time employees receive 20 days of paid annual leave each calendar year. "
        "In addition, employees receive 10 days of paid sick leave. "
        "Leave requests should be submitted at least 7 days in advance for planned vacations. "
        "Unused annual leave may carry over a maximum of 5 days into the next year.",
    )

    pdf.add_page()
    add_heading(pdf, "3. Remote Work")
    add_wrapped_text(
        pdf,
        "Employees may work remotely up to 3 days per week after completing "
        "their first 90 days with the company. Remote work days must be agreed "
        "with the team lead in advance. Employees working remotely must remain "
        "reachable on Slack and email during core collaboration hours. "
        "Fully remote roles are available only for approved positions.",
    )

    add_heading(pdf, "4. Employee Benefits")
    add_wrapped_text(
        pdf,
        "Benefits include health insurance starting on the first day of employment, "
        "a learning stipend of $500 per year, and a home-office stipend of $300 "
        "for employees who work remotely at least 2 days each week. "
        "The company also provides access to an employee assistance program "
        "for mental health support.",
    )

    add_heading(pdf, "5. Performance Reviews")
    add_wrapped_text(
        pdf,
        "Formal performance reviews happen twice each year, in June and December. "
        "Employees set goals with their managers at the start of each review cycle. "
        "Promotion discussions are typically held during the December review window.",
    )

    pdf.output(str(path))


def create_refund_policy_pdf(path: Path) -> None:
    """Create a sample customer refund policy PDF."""
    pdf = SimplePDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.add_page()
    add_heading(pdf, "Customer Refund Policy", size=18)
    add_wrapped_text(
        pdf,
        "This document describes refund eligibility, the refund period, "
        "cancellation rules, and important exceptions for Aspire Demo products.",
    )

    add_heading(pdf, "1. Refund Eligibility")
    add_wrapped_text(
        pdf,
        "Customers may request a refund for digital subscriptions and software "
        "licenses if the product has not been heavily used and the request is "
        "made within the refund period. Physical products must be unused and "
        "returned in original packaging.",
    )

    add_heading(pdf, "2. Refund Period")
    add_wrapped_text(
        pdf,
        "The standard refund period is 30 days from the date of purchase. "
        "Requests submitted after 30 days are not eligible for a refund unless "
        "an exception listed in this policy applies.",
    )

    pdf.add_page()
    add_heading(pdf, "3. Cancellation Rules")
    add_wrapped_text(
        pdf,
        "If you cancel within the first 30 days, you receive a full refund. "
        "If you cancel after 30 days, the remaining subscription period is "
        "non-refundable, but you keep access until the end of the current billing cycle. "
        "Enterprise annual contracts require written cancellation notice "
        "at least 14 days before renewal.",
    )

    add_heading(pdf, "4. International Customers")
    add_wrapped_text(
        pdf,
        "International customers follow the same 30-day refund period. "
        "Refunds are issued in the original payment currency when possible. "
        "Bank or currency-conversion fees charged by the customer's bank "
        "are not reimbursed. Shipping costs for returned physical products "
        "outside the home country are paid by the customer unless the item is defective.",
    )

    add_heading(pdf, "5. Exceptions")
    add_wrapped_text(
        pdf,
        "Exceptions may apply for defective products, accidental duplicate charges, "
        "or service outages lasting more than 72 hours. Gift cards and customized "
        "professional services are non-refundable.",
    )

    pdf.output(str(path))


def create_remote_work_docx(path: Path) -> None:
    """Create a sample remote-work policy DOCX (no page metadata)."""
    doc = DocxDocument()
    doc.add_heading("Remote Work Guidelines", level=1)
    doc.add_paragraph(
        "This document complements the company handbook and focuses on remote "
        "collaboration expectations for Aspire Demo Company employees."
    )

    doc.add_heading("Eligibility", level=2)
    doc.add_paragraph(
        "Remote work is available to employees who have completed onboarding "
        "and received manager approval. Interns may work remotely only one day "
        "per week unless the internship agreement says otherwise."
    )

    doc.add_heading("Communication Expectations", level=2)
    doc.add_paragraph(
        "Remote employees should post their working location and hours in the "
        "team calendar. Video calls are preferred for design reviews and "
        "performance conversations. Response time on urgent Slack messages "
        "should be under 1 hour during core collaboration hours."
    )

    doc.add_heading("Equipment", level=2)
    doc.add_paragraph(
        "The company provides a laptop for all employees. Remote workers may "
        "request a monitor, keyboard, and headset through the IT portal. "
        "Personal devices should not store confidential customer data."
    )

    doc.add_heading("Security", level=2)
    doc.add_paragraph(
        "Employees must use the company VPN when accessing internal systems. "
        "Public Wi-Fi is allowed only with VPN enabled. Sharing passwords "
        "or leaving devices unlocked in shared spaces is not permitted."
    )

    doc.save(str(path))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_company_handbook_pdf(OUTPUT_DIR / "company_handbook.pdf")
    create_refund_policy_pdf(OUTPUT_DIR / "refund_policy.pdf")
    create_remote_work_docx(OUTPUT_DIR / "remote_work_guidelines.docx")
    print("Sample documents created in data/documents/")


if __name__ == "__main__":
    main()
